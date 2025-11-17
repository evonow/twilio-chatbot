"""
Data Processing Script
Loads and processes emails and text messages into the vector database
"""

import os
import json
import re
import csv
import xml.etree.ElementTree as ET
import mailbox
from email import message_from_string
from email.utils import parsedate_to_datetime
from bs4 import BeautifulSoup
from chatbot import ChatbotAgent
from typing import List, Dict, Tuple
import glob

# Optional imports for Word and PDF
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

class DataProcessor:
    def __init__(self):
        """Initialize the data processor with chatbot agent"""
        self.chatbot = ChatbotAgent()
        
    def process_email_file(self, file_path: str, audience: str = None) -> List[Dict]:
        """
        Process a single email file (.eml format)
        
        Args:
            file_path: Path to email file (.eml format)
            audience: Audience label ('sales_reps', 'customers', 'internal', or None)
            
        Returns:
            List of processed email documents
        """
        documents = []
        
        # Try reading as binary first (standard for .eml files)
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
        except Exception as e:
            # Fallback to text mode
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read().encode('utf-8')
            except Exception as e2:
                print(f"Error reading {file_path}: {e2}")
                return documents
        
        # Try to parse as email
        try:
            # Decode to string for message_from_string
            try:
                content_str = content.decode('utf-8')
            except UnicodeDecodeError:
                # Try other encodings
                for encoding in ['latin-1', 'iso-8859-1', 'cp1252']:
                    try:
                        content_str = content.decode(encoding)
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    content_str = content.decode('utf-8', errors='ignore')
            
            msg = message_from_string(content_str)
            subject = msg.get('Subject', 'No Subject')
            
            # Get body - prefer plain text, fallback to HTML
            body = ""
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    if content_type == "text/plain" and not body:
                        try:
                            payload = part.get_payload(decode=True)
                            if payload:
                                body = payload.decode('utf-8', errors='ignore')
                        except Exception:
                            pass
                    elif content_type == "text/html" and not body:
                        try:
                            html_payload = part.get_payload(decode=True)
                            if html_payload:
                                html = html_payload.decode('utf-8', errors='ignore')
                                soup = BeautifulSoup(html, 'html.parser')
                                body = soup.get_text()
                        except Exception:
                            pass
            else:
                try:
                    payload = msg.get_payload(decode=True)
                    if payload:
                        body = payload.decode('utf-8', errors='ignore')
                except Exception:
                    body = msg.get_payload()
            
            # Clean body
            body = self._clean_text(body)
            
            if body:
                from_addr = msg.get('From', 'Unknown')
                to_addr = msg.get('To', 'Unknown')
                cc_addr = msg.get('Cc', '')
                bcc_addr = msg.get('Bcc', '')
                date = msg.get('Date', 'Unknown')
                
                # Build full text with all metadata
                full_text_parts = [
                    f"Email Subject: {subject}",
                    f"From: {from_addr}",
                    f"To: {to_addr}"
                ]
                if cc_addr:
                    full_text_parts.append(f"CC: {cc_addr}")
                if bcc_addr:
                    full_text_parts.append(f"BCC: {bcc_addr}")
                full_text_parts.append(f"Date: {date}")
                full_text_parts.append("")
                full_text_parts.append(body)
                
                full_text = "\n".join(full_text_parts)
                
                # Split into chunks if too long (email embeddings have token limits)
                chunks = self._split_into_chunks(full_text, max_length=2000)
                
                for i, chunk in enumerate(chunks):
                    metadata = {
                        'source': 'email',
                        'file': os.path.basename(file_path),
                        'subject': subject,
                        'from': from_addr,
                        'to': to_addr,
                        'date': date,
                        'chunk_index': i
                    }
                    if cc_addr:
                        metadata['cc'] = cc_addr
                    if bcc_addr:
                        metadata['bcc'] = bcc_addr
                    if audience:
                        metadata['audience'] = audience
                    
                    documents.append({'text': chunk, 'metadata': metadata})
        
        except Exception as e:
            print(f"Error parsing email from {file_path}: {e}")
            import traceback
            traceback.print_exc()
            # Try as plain text fallback
            try:
                content_str = content.decode('utf-8', errors='ignore') if isinstance(content, bytes) else content
                chunks = self._split_into_chunks(self._clean_text(content_str), max_length=2000)
                for i, chunk in enumerate(chunks):
                    metadata = {
                        'source': 'text',
                        'file': os.path.basename(file_path),
                        'chunk_index': i
                    }
                    if audience:
                        metadata['audience'] = audience
                    documents.append({'text': chunk, 'metadata': metadata})
            except Exception as e2:
                print(f"Error processing as plain text: {e2}")
        
        return documents
    
    def process_text_message_file(self, file_path: str, audience: str = None) -> List[Dict]:
        """
        Process a text message file (could be JSON, CSV, XML, or plain text)
        
        Args:
            file_path: Path to text message file
            
        Returns:
            List of processed text message documents
        """
        documents = []
        
        # Try XML first (Android SMS Backup format)
        if file_path.endswith('.xml'):
            return self._process_xml_sms(file_path)
        
        # Try CSV
        if file_path.endswith('.csv'):
            return self._process_csv_sms(file_path)
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            return documents
        
        # Try to parse as JSON first
        try:
            data = json.loads(content)
            
            # Handle different JSON structures
            if isinstance(data, list):
                messages = data
            elif isinstance(data, dict):
                if 'messages' in data:
                    messages = data['messages']
                elif 'texts' in data:
                    messages = data['texts']
                elif 'chats' in data:  # WhatsApp format
                    messages = []
                    for chat in data['chats']:
                        if 'messages' in chat:
                            messages.extend(chat['messages'])
                else:
                    messages = [data]
            else:
                messages = []
            
            for msg in messages:
                if isinstance(msg, dict):
                    text = msg.get('body', msg.get('message', msg.get('text', msg.get('content', ''))))
                    sender = msg.get('from', msg.get('sender', msg.get('phone', msg.get('author', 'Unknown'))))
                    recipient = msg.get('to', msg.get('recipient', msg.get('phone', 'Unknown')))
                    date = msg.get('date', msg.get('timestamp', msg.get('time', 'Unknown')))
                    
                    if text:
                        text = self._clean_text(str(text))
                        full_text = f"SMS/Text Message\nFrom: {sender}\nTo: {recipient}\nDate: {date}\n\n{text}"
                        
                        metadata = {
                            'source': 'sms',
                            'file': os.path.basename(file_path),
                            'from': str(sender),
                            'to': str(recipient),
                            'date': str(date)
                        }
                        documents.append({'text': full_text, 'metadata': metadata})
        
        except json.JSONDecodeError:
            # Not JSON, try as plain text with line-by-line parsing
            lines = content.split('\n')
            current_message = []
            for line in lines:
                line = line.strip()
                if line:
                    # Simple heuristic: messages often start with date/time or phone number
                    if re.match(r'^\d{1,2}[/-]\d{1,2}[/-]\d{2,4}', line) or re.match(r'^\+?\d{10,15}', line):
                        if current_message:
                            text = ' '.join(current_message)
                            if text:
                                metadata = {
                                    'source': 'sms',
                                    'file': os.path.basename(file_path)
                                }
                                if audience:
                                    metadata['audience'] = audience
                                documents.append({'text': self._clean_text(text), 'metadata': metadata})
                            current_message = []
                    current_message.append(line)
            
            # Add last message
            if current_message:
                text = ' '.join(current_message)
                if text:
                    metadata = {
                        'source': 'sms',
                        'file': os.path.basename(file_path)
                    }
                    documents.append({'text': self._clean_text(text), 'metadata': metadata})
        
        return documents
    
    def _process_xml_sms(self, file_path: str) -> List[Dict]:
        """Process Android SMS backup XML file"""
        documents = []
        
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Android SMS backup format
            for sms in root.findall('.//sms'):
                text = sms.get('body', '')
                sender = sms.get('address', 'Unknown')
                date = sms.get('date', 'Unknown')
                msg_type = sms.get('type', '1')  # 1 = received, 2 = sent
                
                if text:
                    text = self._clean_text(text)
                    direction = 'Received' if msg_type == '1' else 'Sent'
                    full_text = f"SMS/Text Message ({direction})\nFrom: {sender}\nDate: {date}\n\n{text}"
                    
                    metadata = {
                        'source': 'sms',
                        'file': os.path.basename(file_path),
                        'from': sender,
                        'date': date,
                        'type': direction
                    }
                    documents.append({'text': full_text, 'metadata': metadata})
        
        except Exception as e:
            print(f"Error parsing XML SMS file {file_path}: {e}")
        
        return documents
    
    def _process_csv_sms(self, file_path: str) -> List[Dict]:
        """Process CSV formatted SMS file"""
        documents = []
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                # Try to detect delimiter
                sample = f.read(1024)
                f.seek(0)
                sniffer = csv.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter
                
                reader = csv.DictReader(f, delimiter=delimiter)
                
                for row in reader:
                    # Common CSV column names
                    text = row.get('body', row.get('message', row.get('text', row.get('content', ''))))
                    sender = row.get('from', row.get('sender', row.get('phone', row.get('address', 'Unknown'))))
                    recipient = row.get('to', row.get('recipient', 'Unknown'))
                    date = row.get('date', row.get('timestamp', row.get('time', 'Unknown')))
                    
                    if text:
                        text = self._clean_text(str(text))
                        full_text = f"SMS/Text Message\nFrom: {sender}\nTo: {recipient}\nDate: {date}\n\n{text}"
                        
                        metadata = {
                            'source': 'sms',
                            'file': os.path.basename(file_path),
                            'from': str(sender),
                            'to': str(recipient),
                            'date': str(date)
                        }
                        documents.append({'text': full_text, 'metadata': metadata})
        
        except Exception as e:
            print(f"Error parsing CSV SMS file {file_path}: {e}")
        
        return documents
    
    def process_mbox_file(self, file_path: str, audience: str = None) -> List[Dict]:
        """
        Process an MBOX file (email archive)
        
        Args:
            file_path: Path to MBOX file
            audience: Audience label ('sales_reps', 'customers', 'internal', or None)
            
        Returns:
            List of processed email documents
        """
        documents = []
        
        try:
            mbox = mailbox.mbox(file_path)
            
            for msg in mbox:
                subject = msg.get('Subject', 'No Subject')
                from_addr = msg.get('From', 'Unknown')
                to_addr = msg.get('To', 'Unknown')
                date = msg.get('Date', 'Unknown')
                
                # Get body
                body = ""
                if msg.is_multipart():
                    for part in msg.walk():
                        content_type = part.get_content_type()
                        if content_type == "text/plain":
                            payload = part.get_payload(decode=True)
                            if payload:
                                body = payload.decode('utf-8', errors='ignore')
                        elif content_type == "text/html" and not body:
                            payload = part.get_payload(decode=True)
                            if payload:
                                html = payload.decode('utf-8', errors='ignore')
                                soup = BeautifulSoup(html, 'html.parser')
                                body = soup.get_text()
                else:
                    payload = msg.get_payload(decode=True)
                    if payload:
                        body = payload.decode('utf-8', errors='ignore')
                
                body = self._clean_text(body)
                
                if body:
                    full_text = f"Email Subject: {subject}\nFrom: {from_addr}\nTo: {to_addr}\nDate: {date}\n\n{body}"
                    chunks = self._split_into_chunks(full_text, max_length=2000)
                    
                    for i, chunk in enumerate(chunks):
                        metadata = {
                            'source': 'email',
                            'file': os.path.basename(file_path),
                            'subject': subject,
                            'from': from_addr,
                            'to': to_addr,
                            'date': date,
                            'chunk_index': i
                        }
                        if audience:
                            metadata['audience'] = audience
                        documents.append({'text': chunk, 'metadata': metadata})
        
        except Exception as e:
            print(f"Error parsing MBOX file {file_path}: {e}")
        
        return documents
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove email headers if present
        text = re.sub(r'^.*?Content-Type:.*?\n', '', text, flags=re.DOTALL)
        # Strip
        text = text.strip()
        return text
    
    def _split_into_chunks(self, text: str, max_length: int = 2000, overlap: int = 200) -> List[str]:
        """
        Split text into overlapping chunks
        
        Args:
            text: Text to split
            max_length: Maximum characters per chunk
            overlap: Overlap between chunks
            
        Returns:
            List of text chunks
        """
        if len(text) <= max_length:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + max_length
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings near the end
                sentence_end = max(
                    text.rfind('.', start, end),
                    text.rfind('!', start, end),
                    text.rfind('?', start, end),
                    text.rfind('\n', start, end)
                )
                if sentence_end > start + max_length // 2:  # Only if we're past halfway
                    end = sentence_end + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
        
        return chunks
    
    def process_directory(self, directory: str, file_pattern: str = "*", audience: str = None):
        """
        Process all files in a directory
        
        Args:
            directory: Directory containing email/text message files
            file_pattern: Glob pattern to match files (default: all files)
            audience: Audience label ('sales_reps', 'customers', 'internal', or None)
        """
        # Find all files
        email_files = glob.glob(os.path.join(directory, f"**/{file_pattern}.eml"), recursive=True)
        mbox_files = glob.glob(os.path.join(directory, f"**/{file_pattern}.mbox"), recursive=True)
        text_files = glob.glob(os.path.join(directory, f"**/{file_pattern}.txt"), recursive=True)
        json_files = glob.glob(os.path.join(directory, f"**/{file_pattern}.json"), recursive=True)
        csv_files = glob.glob(os.path.join(directory, f"**/{file_pattern}.csv"), recursive=True)
        xml_files = glob.glob(os.path.join(directory, f"**/{file_pattern}.xml"), recursive=True)
        
        docx_files = glob.glob(os.path.join(directory, f"**/{file_pattern}.docx"), recursive=True)
        pdf_files = glob.glob(os.path.join(directory, f"**/{file_pattern}.pdf"), recursive=True)
        
        all_files = email_files + mbox_files + text_files + json_files + csv_files + xml_files + docx_files + pdf_files
        
        print(f"Found {len(all_files)} files to process")
        
        total_documents = 0
        
        for file_path in all_files:
            print(f"Processing: {file_path}")
            
            # Determine file type and process accordingly
            if file_path.endswith('.mbox'):
                documents = self.process_mbox_file(file_path, audience=audience)
            elif file_path.endswith('.eml'):
                documents = self.process_email_file(file_path, audience=audience)
            elif file_path.endswith('.docx'):
                documents = self.process_word_document(file_path, audience=audience)
            elif file_path.endswith('.pdf'):
                documents = self.process_pdf_document(file_path, audience=audience)
            elif file_path.endswith('.xml'):
                documents = self.process_text_message_file(file_path, audience=audience)  # XML SMS
            elif file_path.endswith('.csv'):
                # Try SMS first, fall back to email if needed
                if 'sms' in file_path.lower() or 'text' in file_path.lower() or 'message' in file_path.lower():
                    documents = self.process_text_message_file(file_path, audience=audience)
                else:
                    documents = self.process_text_message_file(file_path, audience=audience)  # Try SMS format
            elif 'email' in file_path.lower() or 'mail' in file_path.lower():
                documents = self.process_email_file(file_path, audience=audience)
            elif 'sms' in file_path.lower() or 'text' in file_path.lower() or 'message' in file_path.lower():
                documents = self.process_text_message_file(file_path, audience=audience)
            elif file_path.endswith('.json'):
                documents = self.process_text_message_file(file_path, audience=audience)
            else:
                # Default: try email parsing first
                documents = self.process_email_file(file_path, audience=audience)
            
            # Add documents to knowledge base
            for doc in documents:
                doc_id = f"{os.path.basename(file_path)}_{doc['metadata'].get('chunk_index', 0)}"
                self.chatbot.add_document(
                    text=doc['text'],
                    metadata=doc['metadata'],
                    doc_id=doc_id
                )
                total_documents += 1
        
        print(f"\nProcessing complete! Added {total_documents} documents to knowledge base.")
        return total_documents
    
    def process_google_doc(self, doc_data: Dict, audience: str = None) -> int:
        """
        Process a Google Doc document
        
        Args:
            doc_data: Dictionary with 'content', 'title', 'document_id', etc.
            audience: Audience label ('sales_reps', 'customers', 'internal', or None)
            
        Returns:
            Number of documents added
        """
        content = doc_data.get('content', '')
        title = doc_data.get('title', 'Untitled')
        document_id = doc_data.get('document_id', '')
        
        if not content:
            return 0
        
        # Clean and split into chunks
        cleaned_text = self._clean_text(content)
        chunks = self._split_into_chunks(cleaned_text, max_length=2000)
        
        total_documents = 0
        for i, chunk in enumerate(chunks):
            metadata = {
                'source': 'google_docs',
                'title': title,
                'document_id': document_id,
                'modified_time': doc_data.get('modified_time', ''),
                'created_time': doc_data.get('created_time', ''),
                'chunk_index': i
            }
            if audience:
                metadata['audience'] = audience
            
            doc_id = f"googledoc_{document_id}_{i}"
            self.chatbot.add_document(
                text=chunk,
                metadata=metadata,
                doc_id=doc_id
            )
            total_documents += 1
        
        return total_documents
    
    def process_gitlab_documents(self, documents: List[Dict], audience: str = None) -> int:
        """
        Process GitLab documents (commits, READMEs, release notes)
        
        Args:
            documents: List of dictionaries with 'content' and 'metadata'
            audience: Audience label ('sales_reps', 'customers', 'internal', or None)
            
        Returns:
            Number of documents added
        """
        total_documents = 0
        
        for doc in documents:
            content = doc.get('content', '')
            metadata = doc.get('metadata', {})
            
            if not content:
                continue
            
            # Clean and split into chunks
            cleaned_text = self._clean_text(content)
            chunks = self._split_into_chunks(cleaned_text, max_length=2000)
            
            for i, chunk in enumerate(chunks):
                chunk_metadata = metadata.copy()
                chunk_metadata['chunk_index'] = i
                if audience:
                    chunk_metadata['audience'] = audience
                
                # Create unique doc_id
                source_type = metadata.get('source', 'gitlab')
                if 'commit_id' in metadata:
                    doc_id = f"gitlab_{source_type}_{metadata['commit_id']}_{i}"
                elif 'file_path' in metadata:
                    file_path_safe = metadata['file_path'].replace('/', '_').replace('\\', '_')
                    doc_id = f"gitlab_{source_type}_{file_path_safe}_{i}"
                else:
                    doc_id = f"gitlab_{source_type}_{total_documents}_{i}"
                
                self.chatbot.add_document(
                    text=chunk,
                    metadata=chunk_metadata,
                    doc_id=doc_id
                )
                total_documents += 1
        
        return total_documents
    
    def process_word_document(self, file_path: str, audience: str = None) -> List[Dict]:
        """
        Process a Word document (.docx format)
        
        Args:
            file_path: Path to Word document (.docx format)
            audience: Audience label ('sales_reps', 'customers', 'internal', or None)
            
        Returns:
            List of processed document chunks
        """
        documents = []
        
        if not DOCX_AVAILABLE:
            print("python-docx not available. Install with: pip install python-docx")
            return documents
        
        try:
            doc = Document(file_path)
            
            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n\n".join(text_parts)
            
            if full_text:
                cleaned_text = self._clean_text(full_text)
                chunks = self._split_into_chunks(cleaned_text, max_length=2000)
                
                for i, chunk in enumerate(chunks):
                    metadata = {
                        'source': 'word_document',
                        'file': os.path.basename(file_path),
                        'chunk_index': i
                    }
                    if audience:
                        metadata['audience'] = audience
                    
                    documents.append({'text': chunk, 'metadata': metadata})
        
        except Exception as e:
            print(f"Error processing Word document {file_path}: {e}")
            import traceback
            traceback.print_exc()
        
        return documents
    
    def process_pdf_document(self, file_path: str, audience: str = None) -> List[Dict]:
        """
        Process a PDF document
        
        Args:
            file_path: Path to PDF file
            audience: Audience label ('sales_reps', 'customers', 'internal', or None)
            
        Returns:
            List of processed document chunks
        """
        documents = []
        
        if not PDF_AVAILABLE:
            print("PyPDF2 not available. Install with: pip install PyPDF2")
            return documents
        
        try:
            text_parts = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(f"--- Page {page_num + 1} ---\n{text.strip()}")
                    except Exception as e:
                        print(f"Error extracting text from page {page_num + 1}: {e}")
                        continue
            
            full_text = "\n\n".join(text_parts)
            
            if full_text:
                cleaned_text = self._clean_text(full_text)
                chunks = self._split_into_chunks(cleaned_text, max_length=2000)
                
                for i, chunk in enumerate(chunks):
                    metadata = {
                        'source': 'pdf_document',
                        'file': os.path.basename(file_path),
                        'chunk_index': i
                    }
                    if audience:
                        metadata['audience'] = audience
                    
                    documents.append({'text': chunk, 'metadata': metadata})
        
        except Exception as e:
            print(f"Error processing PDF {file_path}: {e}")
            import traceback
            traceback.print_exc()
        
        return documents

def main():
    """Main function to run data processing"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Process emails and text messages into knowledge base')
    parser.add_argument('directory', help='Directory containing email/text message files')
    parser.add_argument('--pattern', default='*', help='File pattern to match (default: *)')
    
    args = parser.parse_args()
    
    processor = DataProcessor()
    processor.process_directory(args.directory, args.pattern)

if __name__ == '__main__':
    main()


